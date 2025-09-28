"""Lambda handler that renders DOCX and PDF outputs from tailored resume sections."""
from __future__ import annotations

import io
import os
from datetime import datetime, timezone
from typing import Any, Dict, Iterable

import boto3
from docx import Document

s3 = boto3.client("s3")
dynamodb = boto3.resource("dynamodb")

BUCKET_NAME = os.environ["BUCKET_NAME"]
TABLE_NAME = os.environ["TABLE_NAME"]
OUTPUT_PREFIX = os.environ.get("OUTPUT_PREFIX", "generated")


def _flatten_sections(sections: Dict[str, Any]) -> Iterable[str]:
    for section, value in sections.items():
        yield section.upper()
        if isinstance(value, list):
            for entry in value:
                yield f"- {entry}"
        else:
            yield str(value)
        yield ""


def _load_template(template_key: str | None) -> Document:
    if not template_key:
        return Document()
    obj = s3.get_object(Bucket=BUCKET_NAME, Key=template_key)
    return Document(io.BytesIO(obj["Body"].read()))


def _write_docx(document: Document, sections: Dict[str, Any]) -> bytes:
    replaced = False
    for paragraph in document.paragraphs:
        if "{{TAILORED_CONTENT}}" in paragraph.text:
            paragraph.text = ""
            for line in _flatten_sections(sections):
                paragraph.add_run(line)
                paragraph.add_run("\n")
            replaced = True
            break

    if not replaced:
        for line in _flatten_sections(sections):
            paragraph = document.add_paragraph()
            paragraph.add_run(line)
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _to_pdf(text: str) -> bytes:
    safe = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    content_stream = f"BT /F1 11 Tf 72 720 Td ({safe[:2000]}) Tj ET"
    objects = [
        "1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj\n",
        "2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n",
        "3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>endobj\n",
        "4 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj\n",
        f"5 0 obj<< /Length {len(content_stream)} >>stream\n{content_stream}\nendstream\nendobj\n",
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf.extend(obj.encode("utf-8"))
    xref_offset = len(pdf)
    pdf.extend(b"xref\n0 6\n0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010} 00000 n \n".encode("utf-8"))
    pdf.extend(b"trailer<< /Size 6 /Root 1 0 R >>\nstartxref\n")
    pdf.extend(f"{xref_offset}\n".encode("utf-8"))
    pdf.extend(b"%%EOF")
    return bytes(pdf)


def handler(event: Dict[str, Any], _context: Any) -> Dict[str, Any]:
    tenant_id = event["tenantId"]
    job_id = event["jobId"]
    draft = event.get("draft", {})
    template_key = event.get("templateKey")
    sections = draft.get("sections") or {}
    draft_text = draft.get("draftText", "")

    document = _load_template(template_key)
    docx_bytes = _write_docx(document, sections or {"Summary": draft_text})
    pdf_bytes = _to_pdf(draft_text)

    timestamp = datetime.now(timezone.utc).isoformat()
    docx_key = f"{tenant_id}/{OUTPUT_PREFIX}/{job_id}.docx"
    pdf_key = f"{tenant_id}/{OUTPUT_PREFIX}/{job_id}.pdf"

    s3.put_object(Bucket=BUCKET_NAME, Key=docx_key, Body=docx_bytes)
    s3.put_object(Bucket=BUCKET_NAME, Key=pdf_key, Body=pdf_bytes)

    table = dynamodb.Table(TABLE_NAME)
    table.update_item(
        Key={"pk": f"TENANT#{tenant_id}", "sk": f"JOB#{job_id}"},
        UpdateExpression="SET lastRenderedAt = :ts",
        ExpressionAttributeValues={":ts": timestamp},
    )

    return {
        "docxKey": docx_key,
        "pdfKey": pdf_key,
        "completedAt": timestamp,
    }
